from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "initial_report"
ASSET_DIR = ROOT / "outputs" / "initial_report_assets"
DOCX_PATH = OUT_DIR / "复杂航空载体电磁辐射空域特性测量_初期汇报.docx"

FONT_CN = "Microsoft YaHei"
FONT_EN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(96, 96, 96)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def rel(path: str | Path) -> str:
    path = Path(path)
    try:
        return str(path.resolve().relative_to(ROOT)).replace("/", "\\")
    except Exception:
        return str(path).replace("/", "\\")


def read_json(path: str | Path) -> dict:
    with Path(path).open("r", encoding="utf-8") as f:
        return json.load(f)


def read_csv_rows(path: str | Path) -> list[dict[str, str]]:
    with Path(path).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def set_run_font(
    run,
    size: float | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
    italic: bool | None = None,
):
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size: float = 10.5, color: RGBColor | None = None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def add_para(
    doc: Document,
    text: str = "",
    size: float = 10.5,
    bold: bool = False,
    color: RGBColor | None = None,
    after: float = 6,
    before: float = 0,
    line_spacing: float = 1.10,
    align: int | None = None,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.0, color: RGBColor | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_table(
    doc: Document,
    headers: list[str],
    rows: Iterable[Iterable[str]],
    widths: list[float] | None = None,
    font_size: float = 8.8,
):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_text(cell, header, bold=True, size=font_size, color=RGBColor(0, 0, 0))
        set_cell_margins(cell)
        if widths:
            cell.width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_text(cell, str(value), size=font_size)
            set_cell_margins(cell)
            if widths:
                cell.width = Inches(widths[idx])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_numbered(doc: Document, items: Iterable[str]):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{idx}. {item}")
        set_run_font(run, size=10.5)


def add_caption(doc: Document, text: str):
    p = add_para(doc, text, size=9, color=MUTED, after=8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_figure(doc: Document, image_path: str | Path, caption: str, width: float = 6.25):
    image_path = Path(image_path)
    if not image_path.exists():
        add_para(doc, f"图像缺失：{rel(image_path)}", size=9, color=RGBColor(156, 0, 6), italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = FONT_EN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    header_p = section.header.paragraphs[0]
    header_p.text = "复杂航空载体电磁辐射空域特性测量技术比赛 | 初期工程进展汇报"
    set_paragraph_font(header_p, size=8.5, color=MUTED)
    footer_p = section.footer.paragraphs[0]
    footer_p.text = "项目根目录：" + str(ROOT)
    set_paragraph_font(footer_p, size=8, color=MUTED)


def add_cover(doc: Document):
    add_para(doc, "初期工程进展汇报", size=24, bold=True, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "复杂航空载体电磁辐射空域特性测量技术比赛",
        size=15,
        color=MUTED,
        after=18,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    meta_rows = [
        ("报告用途", "帮助队员快速理解现有工程进展，并可直接用于导师/组会初期汇报"),
        ("生成日期", "2026-05-29"),
        ("项目根目录", str(ROOT)),
        ("报告来源", "本机 docs、src、data、outputs、submission 目录扫描；CST 工程截图来自 CST Studio Suite 2025"),
    ]
    add_table(doc, ["项目", "内容"], meta_rows, widths=[1.35, 5.0], font_size=9.2)
    add_para(
        doc,
        "一句话结论：项目已形成“文献依据 - 半球面测量 - CST 标准源/多源样本 - Python 重建与识别 - 结构扰动边界 - 报告/PPT/视频/提交审计”的完整初版工程链。Level 1 required 标准源已完成 2/2，Level 2 full48 样本已完成 48/48，当前识别准确率为 1.000；需要在汇报中保守区分 element-library 证据与完整 full-wave airframe scattering。",
        size=10.5,
        bold=True,
        after=10,
        line_spacing=1.18,
    )
    add_heading(doc, "阅读导引", 2)
    add_numbered(
        doc,
        [
            "先看“总技术路线”，理解为什么先做标准源、再做多源识别、最后做结构遮挡边界。",
            "再看“文献调研”，理解每篇文献被迁移到了哪一段赛题任务。",
            "重点看“CST 建模部分”，这里写清楚每个建模产物对应本机哪个文件夹、为什么这样建模、得到什么信息。",
            "最后看“代码与结果”，理解算法怎么跑、代码入口在哪里、结果说明了什么。",
        ],
    )
    doc.add_page_break()


def add_route_section(doc: Document):
    add_heading(doc, "1. 总技术路线", 1)
    add_para(
        doc,
        "本项目当前采用“先把测量和数据合同标准化，再用低风险 CST 标准源校准链路，随后扩展到多源多状态识别，最后用结构遮挡对照给出安装效应边界”的路线。这样做的核心原因是：复杂航空载体 full-wave 建模成本高、失败定位难；如果先用标准源校准坐标、极化、相位和导出格式，后续复杂样本的可信度更高，也更便于向评委解释。",
        line_spacing=1.15,
    )
    route_rows = [
        ("1", "赛题拆解", "将要求拆成文献调研、2π 测量系统、三维重建、特征辨识和交付物审计。", "outputs\\problem_requirements；outputs\\scorecard"),
        ("2", "文献调研", "用天线/近场测量标准、等效源重构、受限域压缩感知、RF 指纹识别文献约束方案。", "docs\\literature_matrix.md"),
        ("3", "半球面测量", "固定 13 m 上半球 2π、162 个空间测点，覆盖 12m x 10m x 8m 包络。", "outputs\\baseline；outputs\\cst_templates"),
        ("4", "CST Level 1", "短偶极子和半波偶极子标准源，用于坐标、极化、相位和 FarfieldPlot 角域校准。", "outputs\\cst_solver_ready_level1_projects；data\\cst_exports\\level1"),
        ("5", "CST Level 2", "x/y/z element-library 与 48 个多源多状态样本，形成识别数据集。", "outputs\\cst_level2_element_library；data\\cst_exports\\level2"),
        ("6", "Python 算法", "等效源/Tikhonov 重建、最远点少测点选择、SVM/RF 空频极化识别、消融和结构遮挡对照。", "src；outputs\\cst_recognition_level2"),
        ("7", "交付审计", "报告、PPT、视频、代码、CST/data、附录与最终 zip 由审计脚本检查。", "submission；outputs\\completion_audit；outputs\\final_archive"),
    ]
    add_table(doc, ["步骤", "模块", "作用", "本机证据位置"], route_rows, widths=[0.45, 1.25, 3.2, 1.95], font_size=8.5)
    add_figure(
        doc,
        ROOT / "outputs" / "baseline" / "sensor_layout_hemisphere.png",
        "图1 13 m 上半球 2π 测量布局。该图对应 CST 测点表 outputs\\cst_templates\\sensor_layout_hemisphere_for_cst.csv。",
    )
    add_para(
        doc,
        "当前可汇报的主线结论是：测量面已经固定，CST required 标准源与 Level 2 full48 样本链路已完成，算法侧可完成重建、识别、少测点消融和结构扰动边界分析。需要保守说明的是，Level 2 full48 属于 CST-derived element-library 叠加证据，结构遮挡对照是简化迁移模型，不等同于完整复杂载体 full-wave 散射。",
    )


def literature_rows() -> list[tuple[str, str, str, str]]:
    return [
        ("L01", "IEEE 149-2021 天线测量推荐实践", "方向图、测试场、仪器链路、场地评估和极化定义", "测试方案完整性；方向图和远场真值定义"),
        ("L02", "IEEE 1720-2012 近场天线测量推荐实践", "平面/柱面/球面近场测量几何、探头校准、数据交付", "半球面 2π 测量布局；nearfield/farfield CSV 字段"),
        ("L03", "Yaghjian 1986 近场测量综述", "近场测量、采样、探头修正、误差来源", "说明由近场推远场的理论背景"),
        ("L04", "Rodríguez Varela et al. 2020 任意曲面 NF-FF", "任意测量曲面、多级球面波展开和探头修正", "支持半球面/后续不规则布设的高级路线"),
        ("L05", "Kornprobst et al. 2021 表面源重构条件数", "重构面形状影响条件数和精度，球形 Huygens 面较稳", "等效源/Tikhonov 重建；后续贴体 Huygens 面"),
        ("L06", "Malmström et al. 2018 等效天线表示", "等效天线表示会影响平台安装场计算精度", "航空载体安装效应、element-library 边界解释"),
        ("L07", "Valdez/Yuffa/Wakin 2022 受限域压缩感知", "受限球面近远场可用压缩感知减少测点", "2π 半球少测点重建和测点优化"),
        ("L08", "Bangun & Culotta-Lopez 2023 球面测点矩阵优化", "降低感知矩阵互相干性以减少球面测量点", "最远点/均匀性测点选择与随机删点对照"),
        ("L09", "Fuchs et al. 2020 无相位近场测量", "双表面无相位扫描可恢复远场", "幅相数据不稳定时的备用方案"),
        ("L10", "Sarkar & Taaghol 1999 任意几何等效流", "用等效电/磁流和 MoM 做 NF-FF 变换", "复杂测量面条件下的等效源主线"),
        ("L11", "Alvarez/Las-Heras 等三维等效电流重构", "任意三维表面等效电流重构用于诊断和 NF-FF", "后续由规则网格升级为贴近载体外形的重构面"),
        ("L12", "Regué et al. 2001 EMI 源识别与远场预测", "遗传算法识别等效偶极子并预测远场辐射", "多设备等效偶极子/稀疏源定位思路"),
        ("L13", "Gao et al. 2014 磁近场远场预测", "用磁近场扫描和等效定理预测 EMI 远场", "说明近场测量预测远场的工程可信度"),
        ("L14", "Deschrijver et al. 2012 自适应近场扫描", "Kriging/顺序采样减少测量工作量", "主动测点选择、粗扫后局部加密"),
        ("L15", "EMC Europe 2014 幅相同步顺序采样", "同步测量幅度和相位的顺序近场扫描", "主动学习式扫描拓展"),
        ("L16", "2022 混合等效源-粒子群优化", "等效源与粒子群优化提高 NF-FF 稳定性", "若 Tikhonov 不稳定时的对照算法"),
        ("L17", "Soltanieh et al. 2020 RF 指纹综述", "RF 指纹需融合多类特征和分类方法", "空间-频谱-极化联合指纹设计"),
        ("L18", "Liu et al. 2022 雷达发射机机理特征", "频率稳定度、非线性和脉冲包络等多特征优于单特征", "识别模块采用多特征融合而非单一方向图"),
        ("L19", "Gok/Alp/Arikan 2020 真实雷达 SEI", "时间对齐、相干积分和 VMD 提升微弱指纹识别", "多频点/多状态稳健特征提取思想"),
        ("L20", "Gong/Xu/Lei 2020 InfoGAN 无监督 SEI", "非合作少标签场景下的 RF 指纹嵌入", "后续少标签/未知状态识别拓展"),
        ("L21", "Wang et al. 2020 深度复数残差网络 RFFI", "直接处理复数信号减少图像化损失", "后续复数空频向量端到端模型"),
        ("L22", "Jian et al. 2020 大规模 RF 指纹实验", "信道、噪声、训练规模影响识别稳定性", "报告中的噪声/姿态/信道鲁棒性讨论"),
        ("L23", "Man et al. 2021 零样本度量学习 SEI", "用度量学习改善参数变化和零样本泛化", "同一载体状态拉近、不同载体拉远的拓展策略"),
        ("L24", "Jagannath et al. 2022 RF 指纹综合综述", "传统方法、深度学习、数据集和开放问题", "发展趋势和开放挑战补充"),
        ("L25", "Valdez/Folz/Wakin/Gordon 2024 多频稀疏天线测量", "多频稀疏/低秩模型可进一步减少测量量", "Level 2 多频联合采样、频点删减和低秩空频建模"),
    ]


def add_literature_section(doc: Document):
    add_heading(doc, "2. 文献调研及其与赛题的关联", 1)
    add_para(
        doc,
        "目前文献调研不是作为报告装饰，而是被映射到测量规范、近远场重建、少测点优化和特征辨识四条工程主线。项目中的源文件为 docs\\literature_matrix.md、docs\\literature_screening_and_strategy.md 和 docs\\literature_to_algorithm_traceability.md。",
    )
    add_para(
        doc,
        "主引用建议优先使用 L01-L08、L17-L19、L25；其余文献用于备选方案、误差机理、鲁棒性和后续拓展。",
        bold=True,
    )
    add_table(
        doc,
        ["编号", "当前用到的文献", "用到的内容", "关联赛题部分"],
        literature_rows(),
        widths=[0.42, 2.15, 2.35, 1.75],
        font_size=7.7,
    )
    add_heading(doc, "2.1 文献如何约束当前方案", 2)
    add_bullets(
        doc,
        [
            "IEEE 149 和 IEEE 1720 约束了报告中必须写清楚测试场、坐标、极化、探头/测点、远场真值和数据格式，因此项目先生成 CST 测点表和 nearfield/farfield 数据合同。",
            "等效源和表面源重构文献支撑当前的 Tikhonov 等效源反演路线，解释为什么不能只展示 CST 远场图，而要从半球面测量数据重建方向图。",
            "受限域压缩感知和球面感知矩阵优化文献直接对应“2π 空域、尽量少测点”的评分点，因此代码中保留全测点、随机删点、最远点优化删点和鲁棒性扫描。",
            "RF 指纹与雷达 SEI 文献支撑空间-频谱-极化联合特征，不把识别任务简化为单一幅度或单张方向图分类。",
        ],
    )


def add_cst_section(doc: Document):
    add_heading(doc, "3. CST 建模部分", 1)
    add_para(
        doc,
        "CST 建模当前分为三层：Level 1 标准源校准、Level 2 element-library 多源样本、简化结构遮挡对照。它们不是同一个目的：Level 1 用来校准链路，Level 2 用来形成识别数据集，结构遮挡对照用来说明安装/结构效应边界。",
    )
    add_heading(doc, "3.1 本机文件夹对应关系", 2)
    cst_rows = [
        ("测点与导出模板", "outputs\\cst_templates", "半球面 162 测点表、nearfield/farfield 导入导出模板", "给 CST 建模和 Python 校验提供统一数据合同"),
        ("Level 1 API 生成工程", "outputs\\cst_real_level1_projects\\projects", "两类 z 向标准源 .cst 工程", "证明本机 CST API 能自动生成工程"),
        ("Level 1 solver-safe 工程", "outputs\\cst_solver_ready_level1_projects\\projects", "短偶极子和半波偶极子 solver-safe .cst", "用于打开/求解/截图/导出 Level 1 required 数据"),
        ("Level 1 VBA 建模历史", "outputs\\cst_solver_ready_level1_projects\\vba_history", "单位、边界、偶极子、端口、监视器和 162 探针的 VBA 记录", "可复核建模过程，必要时手工重放"),
        ("Level 1 真实导出", "data\\cst_exports\\level1", "required 两个标准源 nearfield/farfield CSV 及合并表", "供合并审计、角域校准和重建使用"),
        ("Level 1 角域校准", "outputs\\cst_level1_angular_calibration", "两类标准源的角域拟合结果和对比图", "验证 FarfieldPlot-derived 方向图一致性"),
        ("Level 2 element 工程", "outputs\\cst_level2_element_library\\projects", "x/y/z 三方向短偶极子、5 个频点的 .cst", "避免直接求解大尺度多源机体导致网格过大"),
        ("Level 2 叠加导出", "outputs\\cst_level2_superposed_export", "按 48 个样本和 240 个 sample-frequency 生成数据", "形成识别样本和消融实验数据源"),
        ("Level 2 真实导出表", "data\\cst_exports\\level2", "48 样本合并 nearfield/farfield CSV", "供 SVM/RF 识别和少测点/频点消融"),
        ("结构遮挡对照", "outputs\\cst_structure_comparison", "简化机身/机翼/尾翼遮挡迁移后的方向图和识别结果", "量化安装结构影响边界，避免夸大 element-library"),
        ("提交包 CST 目录", "submission\\05_cst", "上述工程、任务包、审计证据的提交副本", "用于最终材料整理和人工核查"),
    ]
    add_table(doc, ["产物", "本机文件夹", "内容", "用途"], cst_rows, widths=[1.1, 2.0, 2.0, 1.55], font_size=7.8)
    add_heading(doc, "3.2 重要 CST 建模截图", 2)
    add_figure(
        doc,
        ASSET_DIR / "cst_halfwave_foreground_full.png",
        "图2 CST Studio Suite 2025 中打开的 Level 1 半波偶极子 z 向工程。对应工程：outputs\\cst_solver_ready_level1_projects\\projects\\CST_L1_halfwave_dipole_z_1p2G.cst。",
    )
    add_para(
        doc,
        "该模型采用 PEC 圆柱上下两臂加中心离散端口，工作频点 1.2 GHz，开放边界，设置 E-field monitor 和 Farfield monitor，并插入 162 个半球面 efarfield 探针。这样建模的目的不是模拟复杂机体，而是先用标准源确认坐标轴、端口、极化、相位和远场导出是否被 Python 正确读取。",
    )
    add_figure(
        doc,
        ASSET_DIR / "cst_level2_element_x_foreground.png",
        "图3 CST Studio Suite 2025 中打开的 Level 2 x 向短偶极子 element-library 工程。对应工程：outputs\\cst_level2_element_library\\projects\\CST_L2_element_short_dipole_x_5freq.cst。",
    )
    add_para(
        doc,
        "Level 2 没有直接把 48 个样本的所有辐射源一次性放入完整大尺度空域求解，而是先求解 x/y/z 三个短偶极子 element，再按样本 manifest 中的位置、方向、幅度和相位做线性叠加。这样可以显著降低本机 CST 网格和求解压力，并且让每个样本的数据生成可审计、可复现。",
    )
    add_heading(doc, "3.3 CST 建模得到的信息和用途", 2)
    level1 = read_json(ROOT / "outputs" / "cst_level1_angular_calibration" / "angular_calibration_summary.json")
    level2 = read_json(ROOT / "outputs" / "cst_level2_merge_report" / "level2_merge_summary.json")
    structure = read_json(ROOT / "outputs" / "cst_structure_comparison" / "structure_comparison_summary.json")
    cst_result_rows = [
        ("Level 1 required", "2/2 标准源完成；required_complete=true", "校准坐标、相位、极化和 FarfieldPlot 导出链路"),
        ("Level 1 角域校准", f"max NMSE={level1['max_nmse']:.2e}；min corr={level1['min_correlation']:.5f}；max main-lobe error={level1['max_main_lobe_error_deg']:.1f} deg", "说明标准源角域方向图和 CST 远场高度一致"),
        ("Level 2 full48", f"48/48 样本；240/240 sample-frequency；nearfield {level2['merged_nearfield_rows']} 行；farfield {level2['merged_farfield_rows']} 行", "作为空间-频谱-极化识别数据集"),
        ("结构遮挡边界", f"mean shadow={structure['mean_shadow_db']:.2f} dB；P95={structure['p95_shadow_db']:.2f} dB；cross-domain accuracy={structure['cross_domain_accuracy']:.3f}", "说明安装/结构扰动会改变方向图，但当前识别特征在简化扰动下仍稳健"),
    ]
    add_table(doc, ["模块", "得到的信息", "用来干什么"], cst_result_rows, widths=[1.3, 3.05, 2.15], font_size=8.4)
    add_figure(
        doc,
        ROOT / "outputs" / "cst_level1_angular_calibration" / "L1_halfwave_dipole_z_1p2G" / "angular_farfield_compare.png",
        "图4 Level 1 半波偶极子 FarfieldPlot-derived 角域校准对比图。该图用于说明标准源链路的方向图一致性。",
    )
    add_figure(
        doc,
        ROOT / "outputs" / "cst_structure_comparison" / "plots" / "L2_radar_top_000_1200MHz_structure_compare.png",
        "图5 Level 2 简化结构遮挡对照示例。该图用于说明结构/安装扰动会改变方向图，是边界说明而不是 full-wave 机体散射结论。",
    )


def add_code_section(doc: Document):
    add_heading(doc, "4. 代码部分：算法、职责、结果和含义", 1)
    add_para(
        doc,
        "代码集中在 src 目录。可以把它理解成六类：基础电磁算法、CST 数据接口、CST 自动化/任务包、合并重建识别、结构对照、报告审计交付。核心算法不是单一分类器，而是“物理重建 + 少测点优化 + 空频极化识别 + 审计闭环”。",
    )
    code_rows = [
        ("src\\em_core.py", "半球面测点、球坐标极化基、等效源网格、偶极子场、测量矩阵、Tikhonov、远场和指标", "支撑近场到远场重建主算法"),
        ("src\\run_baseline.py", "生成 162 测点半球面、合成场、全/随机/优化删点重建、基础识别", "验证算法链路可运行，输出 baseline 图表"),
        ("src\\cst_io.py", "读取 CST CSV、校验字段、Ex/Ey/Ez 到 theta/phi 投影、构造测量向量", "保证 CST 导出能被 Python 稳定接入"),
        ("src\\check_cst_export.py", "对 nearfield/farfield 做完整性和数值校验", "作为真实 CST 导出进入算法前的 gate"),
        ("src\\run_cst_reconstruction.py", "单样本等效源/Tikhonov 反演和方向图对比", "从 CST nearfield 重建远场并计算误差"),
        ("src\\run_cst_level1_angular_calibration.py", "用角域基函数和正则化拟合 FarfieldPlot 方向图", "给 Level 1 标准源提供稳定的角域一致性证据"),
        ("src\\run_cst_recognition.py", "提取归一化幅度、相位余弦/正弦、极化功率比、空间功率摘要和频谱特征；训练 SVM/RF", "对应赛题特征辨识准确率要求"),
        ("src\\run_cst_recognition_ablation.py", "按最远点删减测点，按中心频点删减频率，重复训练识别器", "证明少测点/少频点下识别链路稳定性"),
        ("src\\run_cst_structure_comparison.py", "对 Level 2 场施加简化机身/机翼/尾翼遮挡迁移，统计方向图变化和跨域识别", "量化结构/安装效应边界"),
        ("src\\build_*", "scorecard、problem requirements、completion audit、submission index、final archive、报告/PPT/进展", "把技术结果转为可汇报、可提交、可审计材料"),
    ]
    add_table(doc, ["代码", "做什么", "用途"], code_rows, widths=[1.55, 3.35, 1.65], font_size=8.0)
    add_heading(doc, "4.1 使用的主要算法", 2)
    add_bullets(
        doc,
        [
            "等效源反演：把半球面近场测量写成 y = G j + n，通过 Tikhonov 正则求解等效源系数，再外推远场方向图。",
            "少测点优化：用 farthest point subset 选择几何分散的测点子集，与随机删点和全测点对比，用 NMSE、相关系数和主瓣误差评估。",
            "角域校准：对 CST FarfieldPlot-derived 远场做正则化角域基函数拟合，作为 Level 1 标准源的稳定一致性证据。",
            "特征识别：对每个样本提取归一化 log 幅度、相对相位、极化功率比、空间能量分布摘要、频点间功率变化，再用 SVM RBF 和 Random Forest 分类。",
            "结构边界模型：对 element-library 场施加简化航空载体遮挡迁移，统计 shadow dB、方向图相关性、跨域识别准确率。",
        ],
    )
    add_heading(doc, "4.2 当前关键结果", 2)
    rec_rows = read_csv_rows(ROOT / "outputs" / "baseline" / "reconstruction_metrics.csv")
    opt50 = next(row for row in rec_rows if row["experiment"] == "optimized_50")
    full = next(row for row in rec_rows if row["experiment"] == "full_100")
    recog = read_json(ROOT / "outputs" / "cst_recognition_level2" / "cst_recognition_metrics.json")
    ablation = read_json(ROOT / "outputs" / "cst_recognition_level2_ablation" / "recognition_ablation_summary.json")
    structure = read_json(ROOT / "outputs" / "cst_structure_comparison" / "structure_comparison_summary.json")
    result_rows = [
        ("baseline 全测点重建", f"162 点/324 通道；NMSE={float(full['nmse']):.2e}；corr={float(full['correlation']):.5f}", "全测点链路可作为上限参考"),
        ("baseline 50% 优化删点", f"81 点/162 通道；NMSE={float(opt50['nmse']):.2e}；corr={float(opt50['correlation']):.5f}", "少测点后仍能保持较高方向图相关性"),
        ("Level 1 角域校准", "max NMSE=8.41e-5；min corr=0.99988；主瓣误差 0 deg", "标准源 FarfieldPlot-derived 方向图一致性很高"),
        ("Level 2 识别", f"{recog['feature_metadata']['sample_count']} 样本；{recog['feature_metadata']['feature_count']} 维特征；best={recog['recognition']['best_model']}；accuracy={recog['recognition']['best_accuracy']:.3f}", "满足并超过 85% 特征辨识要求"),
        ("Level 2 消融", f"{len(ablation['ablation_cases'])} 组测点/频点组合；最低 accuracy={min(float(x['best_accuracy']) for x in ablation['ablation_cases']):.3f}", "说明 50% 测点、1/3/5 频点组合下识别仍稳定"),
        ("结构遮挡对照", f"mean shadow={structure['mean_shadow_db']:.2f} dB；mean pattern corr={structure['mean_pattern_correlation']:.3f}；cross-domain accuracy={structure['cross_domain_accuracy']:.3f}", "说明结构扰动影响方向图，但当前识别特征仍有鲁棒性"),
    ]
    add_table(doc, ["结果", "数值", "说明"], result_rows, widths=[1.45, 2.85, 2.2], font_size=8.4)
    add_figure(
        doc,
        ROOT / "outputs" / "cst_recognition_level2" / "cst_recognition_confusion_matrix.png",
        "图6 Level 2 full48 空间-频谱-极化识别混淆矩阵。当前测试集分类准确率为 1.000。",
    )
    add_figure(
        doc,
        ROOT / "outputs" / "cst_recognition_level2_ablation" / "recognition_ablation_accuracy.png",
        "图7 Level 2 测点/频点消融结果。该图用于说明少测点和少频点条件下识别链路的稳定性。",
    )


def add_report_section(doc: Document):
    add_heading(doc, "5. 当前工程进展汇报口径", 1)
    add_heading(doc, "5.1 可以直接讲的结论", 2)
    add_bullets(
        doc,
        [
            "项目已经完成从文献依据、测量布局、CST 工程、数据导出、Python 算法到报告/PPT/视频/压缩包的初版闭环。",
            "测量系统采用 13 m 上半球 2π 布局，共 162 个空间测点，满足 12m x 10m x 8m 被测对象包络的外部采样需求。",
            "Level 1 required 标准源完成 2/2，FarfieldPlot-derived 角域校准结果很稳定：最大 NMSE 约 8.41e-5，最小相关系数约 0.99988。",
            "Level 2 full48 样本完成 48/48，形成 240 个 sample-frequency，SVM/RF 识别准确率为 1.000。",
            "结构遮挡对照给出 mean shadow 约 3.06 dB、P95 shadow 约 6.63 dB，跨域识别准确率仍为 1.000。",
        ],
    )
    add_heading(doc, "5.2 需要保守说明的边界", 2)
    add_bullets(
        doc,
        [
            "Level 1 高精度结论是 FarfieldPlot-derived 角域一致性，不应夸大为 full-wave 近场 monitor 等效源反演已经完全解决。",
            "Level 2 full48 属于 CST-derived element-library 线性叠加样本，不是完整复杂航空载体 full-wave airframe scattering。",
            "结构遮挡对照是透明可复现的简化安装/遮挡迁移模型，用于边界分析，不应替代真正的复杂机体全波仿真。",
            "当前视频是 PowerPoint 自动计时静音版，正式提交前应人工完整播放，并按赛事要求决定是否替换为带讲解录屏。",
        ],
    )
    add_heading(doc, "5.3 汇报建议顺序", 2)
    add_numbered(
        doc,
        [
            "先讲赛题要求拆解：调研、2π 测量、三维重建、特征辨识和交付物。",
            "讲文献如何决定方案：测量标准、等效源、受限域压缩感知、多频测量和 RF 指纹。",
            "讲 CST 为什么分层：标准源先校准，element-library 再扩展多源，结构对照负责边界。",
            "讲代码如何闭环：CST CSV 进入 Python，完成重建、识别、消融、结构扰动和审计。",
            "讲当前结果和边界：数据已经跑通且指标好，但 full-wave 复杂机体散射仍是后续增强项。",
        ],
    )
    add_heading(doc, "5.4 后续优先事项", 2)
    add_table(
        doc,
        ["优先级", "任务", "目的"],
        [
            ("1", "人工播放并确认 submission\\03_video\\demo_video.mp4", "避免正式报送时视频黑屏、卡顿或不符合讲解要求"),
            ("2", "补齐 submission\\00_admin 中队伍、学校、申报人、联系方式等人工信息", "完成正式提交材料"),
            ("3", "若时间允许，增加复杂载体 full-wave airframe 对照", "提升结构散射证据强度"),
            ("4", "如替换视频或改提交内容，重跑 build_submission_draft.py、build_final_archive.py、build_completion_audit.py", "保持 completion_proven 和 zip 清单一致"),
        ],
        widths=[0.7, 3.4, 2.4],
        font_size=8.6,
    )


def add_appendix(doc: Document):
    add_heading(doc, "附录：本报告引用的关键本机文件", 1)
    rows = [
        ("项目总入口", "README.md"),
        ("当前工作详细讲解", "docs\\current_work_detailed_explanation.md"),
        ("文献矩阵", "docs\\literature_matrix.md"),
        ("文献到算法证据链", "docs\\literature_to_algorithm_traceability.md"),
        ("项目进展报告", "docs\\project_progress_report.md"),
        ("项目文件索引", "docs\\project_file_index.md"),
        ("Level 1 工程 manifest", "outputs\\cst_solver_ready_level1_projects\\cst_level1_project_manifest.csv"),
        ("Level 2 工程 manifest", "outputs\\cst_level2_element_library\\element_project_manifest.csv"),
        ("Level 1 角域校准摘要", "outputs\\cst_level1_angular_calibration\\angular_calibration_summary.json"),
        ("Level 2 识别指标", "outputs\\cst_recognition_level2\\cst_recognition_metrics.json"),
        ("Level 2 消融摘要", "outputs\\cst_recognition_level2_ablation\\recognition_ablation_summary.json"),
        ("结构遮挡摘要", "outputs\\cst_structure_comparison\\structure_comparison_summary.json"),
        ("完成度审计", "outputs\\completion_audit\\completion_audit.md"),
        ("最终提交草稿", "submission"),
    ]
    add_table(doc, ["用途", "路径"], rows, widths=[1.75, 4.75], font_size=8.5)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_route_section(doc)
    add_literature_section(doc)
    add_cst_section(doc)
    add_code_section(doc)
    add_report_section(doc)
    add_appendix(doc)
    doc.core_properties.title = "复杂航空载体电磁辐射空域特性测量技术比赛初期汇报"
    doc.core_properties.subject = "初期工程进展、文献、CST 建模、代码算法与结果"
    doc.core_properties.author = "Codex"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
