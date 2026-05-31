# -*- coding: utf-8 -*-
from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "initial_report_focused"
DOCX_PATH = OUT_DIR / "复杂航空载体电磁辐射空域特性测量_初期汇报_聚焦版.docx"

FONT_CN = "Microsoft YaHei"
FONT_EN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 32, 32)
MUTED = RGBColor(96, 96, 96)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def read_json(path: str | Path) -> dict:
    with Path(path).open("r", encoding="utf-8") as f:
        return json.load(f)


def read_csv_rows(path: str | Path) -> list[dict[str, str]]:
    with Path(path).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def rel(path: str | Path) -> str:
    try:
        return str(Path(path).resolve().relative_to(ROOT)).replace("/", "\\")
    except Exception:
        return str(path).replace("/", "\\")


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def set_paragraph_font(paragraph, size=10.5, color=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold=False, size=9.0, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color or INK)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_table_widths(table, widths: list[float]):
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[str]], widths: list[float], font_size=8.8):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_text(cell, header, bold=True, size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value), size=font_size)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_para(
    doc: Document,
    text: str = "",
    size=10.5,
    bold=False,
    color=None,
    after=6,
    before=0,
    line_spacing=1.12,
    align=None,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color or INK)
    return p


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=11.5, bold=True, color=DARK_BLUE)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def add_callout(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.25)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=130, start=150, bottom=130, end=150)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.1
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.0, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_caption(doc: Document, text: str):
    p = add_para(doc, text, size=9.0, color=MUTED, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p


def add_figure(doc: Document, image_path: str | Path, caption: str, width=6.25):
    path = Path(image_path)
    if not path.exists():
        add_para(doc, f"图像缺失：{rel(path)}", size=9, color=RGBColor(156, 0, 6), italic=True)
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_ref(doc: Document, text: str):
    p = add_para(doc, "参考依据：" + text, size=9.0, color=MUTED, after=6)
    p.paragraph_format.left_indent = Inches(0.18)
    return p


def configure_document(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = styles[name]
        style.font.name = FONT_EN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    header = section.header.paragraphs[0]
    header.text = "复杂航空载体电磁辐射空域特性测量技术比赛 | 初期汇报聚焦版"
    set_paragraph_font(header, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "项目根目录：" + str(ROOT)
    set_paragraph_font(footer, size=8.0, color=MUTED)


def add_cover(doc: Document):
    add_para(doc, "初期汇报聚焦版", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(
        doc,
        "CST 建模设置、建模结果与代码算法说明",
        size=15,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=18,
    )
    add_table(
        doc,
        ["项目", "内容"],
        [
            ("汇报目标", "只讲 CST 建模与代码算法两大部分，帮助快速理解现阶段工程进展。"),
            ("写法调整", "文献不再单独成章；在对应技术内容后标注参考依据。"),
            ("图片来源", "插入本机 CST 截图、半球测点图、远场重建图、识别混淆矩阵和消融结果图。"),
            ("生成时间", "2026-05-29"),
            ("项目根目录", str(ROOT)),
        ],
        widths=[1.35, 4.9],
        font_size=9.3,
    )
    add_callout(
        doc,
        "阅读导引",
        "这份报告按一个问题展开：我们在 CST 里建了什么，得到哪些电磁数据；这些数据进入 Python 后，算法如何完成远场重建、少测点验证和空间-频谱识别。",
    )


def fmt_float(value, digits=4):
    return f"{float(value):.{digits}g}"


def add_cst_section(doc: Document, data: dict):
    add_heading(doc, "一、CST 建模部分", 1)
    add_callout(
        doc,
        "本部分一句话",
        "CST 的作用是把抽象的辐射源和测量布局变成可检查、可导出的电磁场数据；当前工程采用“标准源校准 -> 多源基元库 -> 结构扰动对照”的分层建模方式。",
    )

    add_heading(doc, "1.1 为什么 CST 建模要分三层", 2)
    add_para(
        doc,
        "复杂航空载体上的真实辐射源数量多、位置复杂、工作频率不同，如果一开始就做完整机体 full-wave 求解，模型规模和网格规模都会很大，调试也很困难。当前工程先用简单且可解释的标准源校准测量链路，再用三方向短偶极子基元库扩展到多源多状态样本，最后用简化结构遮挡模型评估安装结构对方向图和识别的影响。",
    )
    add_table(
        doc,
        ["层级", "CST 建模内容", "主要回答的问题", "当前产物"],
        [
            (
                "Level 1 标准源",
                "1.2 GHz 短偶极子、半波偶极子；每个工程插入 162 个半球面 efarfield 探针。",
                "坐标、极化、CST 导出和 Python 读取是否一致？典型源远场方向图能否对上？",
                "2/2 required cases 完成；角域 max NMSE 8.41e-5。",
            ),
            (
                "Level 2 基元库",
                "x/y/z 三方向短偶极子，每个基元覆盖 900-1500 MHz 五个频点。",
                "如何在本机算力范围内构造多源、多频、多状态识别数据？",
                "3/3 基元工程完成；叠加形成 48 个样本、240 个 sample-frequency。",
            ),
            (
                "结构对照",
                "对 Level 2 场数据施加可复现的机身/机翼/尾翼遮挡转移函数。",
                "结构安装和遮挡会把方向图改多少？识别链路在结构扰动下是否仍稳定？",
                "mean shadow 3.06 dB；cross-domain accuracy 1.000。",
            ),
        ],
        widths=[1.0, 1.85, 2.05, 1.35],
        font_size=8.4,
    )
    add_ref(doc, "L01 IEEE 149-2021；L02 IEEE 1720-2012；L03 Yaghjian 近场测量综述；L06 Malmstrom 等效天线表示。")

    add_heading(doc, "1.2 测量布局如何体现在 CST 中", 2)
    add_para(
        doc,
        "测量方案的源头在 Python：代码统一生成 13 m 半径、覆盖上半空间 2π 立体角的 162 个半球面测点，并给每个测点固定编号和空间坐标。这样做的好处是避免手工布点误差，也方便后续把同一套测点用于 CST、重建和识别。",
    )
    add_para(
        doc,
        "在 Level 1 CST 工程中，这 162 个点已经作为 efarfield 探针写入真实 CST 工程文件；也就是说，Level 1 不是只在 Python 后处理里假想测点，而是已经把测点布设写进 CST 项目。Level 2 为了控制网格规模，采用基元库和后处理叠加方式，输出数据仍按同一套 162 点测量布局组织。",
    )
    add_table(
        doc,
        ["参数", "当前设置", "直白解释"],
        [
            ("测量面", "上半球 2π 空间", "只覆盖被测体上方半空间，满足赛题半球面/半柱面 2π 要求。"),
            ("半径", "13 m", "大于 12m x 10m x 8m 包络的外部测量半径，给目标留出空间。"),
            ("测点数", "162 个空间点", "不是高密度全波长采样，而是服务于等效源反演和识别的稀疏工程布局。"),
            ("极化/分量", "Level 1: efarfield；Level 2 CSV: Ex/Ey/Ez，可转 theta/phi", "既保留三维场信息，也能转换为球面极化特征。"),
            ("频点", "Level 2: 900/1050/1200/1350/1500 MHz", "用五个离散频点表达宽频段采集能力。"),
        ],
        widths=[1.15, 1.8, 3.3],
        font_size=8.8,
    )
    add_figure(
        doc,
        ROOT / "outputs" / "baseline" / "sensor_layout_hemisphere.png",
        "图 1  13 m 上半球 162 测点布局。该图对应代码生成的测量点坐标，也是 CST/数据表统一使用的空间采样框架。",
        width=5.15,
    )
    add_ref(doc, "L02 IEEE 1720 支撑球面/半球面近场测量几何；L07 受限域压缩感知支撑 2π 受限空域下的少测点思路。")

    add_heading(doc, "1.3 Level 1 标准源：先校准链路，再谈复杂载体", 2)
    add_para(
        doc,
        "Level 1 使用典型偶极子作为标准源。原因很简单：偶极子的辐射方向图有明确物理规律，容易判断坐标轴、极化方向、主瓣位置和远场形态是否正确。如果标准源都对不上，直接做复杂载体没有意义。",
    )
    add_para(
        doc,
        "当前 Level 1 建了两个 required 标准源：z 向短偶极子和 z 向半波偶极子，频率均为 1.2 GHz。CST 工程由真实 CST API 生成，工程文件、VBA history、输入快照和自动化日志都保存在项目目录中，可以打开检查。",
    )
    add_table(
        doc,
        ["内容", "位置"],
        [
            ("Level 1 solver-ready CST 工程", rel(ROOT / "outputs" / "cst_solver_ready_level1_projects" / "projects")),
            ("Level 1 自动化日志和摘要", rel(ROOT / "outputs" / "cst_solver_ready_level1_projects")),
            ("Level 1 合并后的 nearfield/farfield CSV", rel(ROOT / "data" / "cst_exports" / "level1")),
            ("Level 1 角域校准结果", rel(ROOT / "outputs" / "cst_level1_angular_calibration")),
        ],
        widths=[2.05, 4.2],
        font_size=8.6,
    )
    add_figure(
        doc,
        ROOT / "outputs" / "initial_report_assets" / "cst_halfwave_foreground_full.png",
        "图 2  CST 中打开的 Level 1 半波偶极子工程截图。它用于检查标准源几何、坐标方向和 CST 工程是否真实可打开。",
        width=6.25,
    )
    l1_merge = data["l1_merge"]
    l1_ang = data["l1_ang"]
    add_table(
        doc,
        ["结果项", "当前值", "说明"],
        [
            ("required cases", f"{l1_merge['required_complete_cases']}/{l1_merge['required_cases']}", "赛题/工程要求的 Level 1 标准源已经齐套。"),
            ("merged nearfield rows", str(l1_merge["merged_nearfield_rows"]), "近场采样数据已进入统一 CSV。"),
            ("merged farfield rows", str(l1_merge["merged_farfield_rows"]), "远场角域数据已进入统一 CSV。"),
            ("angular max NMSE", fmt_float(l1_ang["max_nmse"], 4), "角域拟合误差很小。"),
            ("angular min correlation", fmt_float(l1_ang["min_correlation"], 6), "CST 远场方向图和校准模型高度相关。"),
            ("main-lobe error", f"{l1_ang['max_main_lobe_error_deg']} deg", "主瓣方向没有偏移。"),
        ],
        widths=[1.75, 1.4, 3.1],
        font_size=8.8,
    )
    add_figure(
        doc,
        ROOT / "outputs" / "cst_level1_angular_calibration" / "L1_halfwave_dipole_z_1p2G" / "angular_farfield_compare.png",
        "图 3  Level 1 标准源角域校准示例。曲线重合度高，说明坐标、极化和 FarfieldPlot 导出链路基本可信。",
        width=5.8,
    )
    add_para(
        doc,
        "这部分的用途是作为工程地基：证明 CST 工程、测点、导出格式和 Python 读取没有明显方向性错误。需要保守说明的是，这里的高精度结论是 FarfieldPlot-derived 角域一致性，不等于已经完成复杂机体 full-wave 近场反演。",
    )
    add_ref(doc, "L03 Yaghjian 说明近场测量到远场验证的基本逻辑；L05 Kornprobst 支撑等效源/表面源反演需要关注条件数和正则化。")

    add_heading(doc, "1.4 Level 2 基元库：用可控方式构造多源多状态样本", 2)
    add_para(
        doc,
        "Level 2 的目标不是再证明单个偶极子，而是构造能够训练和验证识别算法的多源、多频、多状态样本。直接把完整 12m 级复杂载体和所有源都放进一个 CST full-wave 空气域中，本机求解成本很高。因此当前采用 element-library 路线：先在 CST 中分别建立 x、y、z 三方向短偶极子基元，每个基元覆盖五个频点，再在 Python 后处理中按源位置、幅度和相位叠加为 48 个样本。",
    )
    add_table(
        doc,
        ["建模设置", "当前值"],
        [
            ("基元方向", "x、y、z 三方向短偶极子"),
            ("频率", "900、1050、1200、1350、1500 MHz"),
            ("CST 工程数量", "3 个基元工程，均由真实 CST API 生成"),
            ("样本类别", "comm_pair、mixed_avionics、multi_state_on、radar_top"),
            ("样本数量", "48 个样本，每类 12 个"),
            ("sample-frequency", "48 x 5 = 240 组"),
            ("合并数据规模", "nearfield 116640 行，farfield 164160 行"),
        ],
        widths=[1.8, 4.45],
        font_size=8.8,
    )
    add_figure(
        doc,
        ROOT / "outputs" / "initial_report_assets" / "cst_level2_element_x_foreground.png",
        "图 4  CST 中打开的 Level 2 x 向短偶极子基元工程截图。三方向基元库用于后续合成多源、多状态样本。",
        width=6.25,
    )
    add_para(
        doc,
        "这样建模的好处是可解释、可复现、算力压力小。每个复杂样本都可以追溯到若干个等效辐射基元的空间位置、方向、幅度和相位；如果识别结果异常，也能回头检查是哪一类源组合或哪一个频点导致问题。",
    )
    add_ref(doc, "L06 Malmstrom 等效天线表示支持用等效源降低平台安装问题复杂度；L25 多频稀疏测量支持多频联合建模。")

    add_heading(doc, "1.5 结构对照：说明当前证据边界", 2)
    add_para(
        doc,
        "当前工程还做了一个简化结构遮挡对照。它不是完整机体 full-wave CST 求解，而是在 Level 2 场数据基础上加入可复现的机身、机翼、尾翼遮挡转移函数，然后观察方向图和识别结果如何变化。",
    )
    s = data["structure"]
    add_table(
        doc,
        ["指标", "结果", "直白解释"],
        [
            ("mean shadow", f"{s['mean_shadow_db']:.2f} dB", "平均遮挡会带来约 3 dB 量级的方向图变化。"),
            ("P95 shadow", f"{s['p95_shadow_db']:.2f} dB", "大多数方向变化低于约 6.6 dB。"),
            ("max shadow", f"{s['max_shadow_db']:.2f} dB", "个别方向遮挡影响较强。"),
            ("mean pattern correlation", f"{s['mean_pattern_correlation']:.3f}", "结构会显著改变方向图形态。"),
            ("cross-domain accuracy", f"{s['cross_domain_accuracy']:.3f}", "在当前样本上，识别链路对该结构扰动仍稳定。"),
        ],
        widths=[1.7, 1.3, 3.25],
        font_size=8.8,
    )
    add_figure(
        doc,
        ROOT / "outputs" / "cst_structure_comparison" / "plots" / "L2_comm_pair_000_1200MHz_structure_compare.png",
        "图 5  简化结构遮挡前后方向图对比。它用于说明载体结构会改变辐射空间分布，也用于给当前结果划清边界。",
        width=5.9,
    )
    add_para(
        doc,
        "汇报时建议这样说：我们已经完成了结构影响的有界对照，证明算法链路可以承受一定安装/遮挡扰动；但完整复杂航空载体 full-wave 结构散射仍是后续增强项，不应把当前 element-library 叠加结果说成完整全机电磁仿真。",
    )
    add_ref(doc, "L05 Kornprobst 和 L06 Malmstrom 都提示复杂结构安装效应会影响等效源/方向图结果，需要单独说明模型边界。")


def add_code_section(doc: Document, data: dict):
    add_heading(doc, "二、代码算法部分", 1)
    add_callout(
        doc,
        "本部分一句话",
        "代码把 CST/仿真数据变成三件事：远场重建、少测点验证、空间-频谱识别。核心算法是等效源正则化反演、几何均匀测点抽取、空间-频谱特征工程和 SVM/RF 分类。",
    )

    add_heading(doc, "2.1 代码总流程", 2)
    add_para(
        doc,
        "整个代码链路可以按数据流理解：先生成半球测点和等效源网格，再从 CST 或仿真模型得到测量数据；然后一条支路做等效源反演和远场重建，另一条支路做空间-频谱特征识别；最后用消融实验和结构对照验证少测点、少频点和结构扰动下的稳定性。",
    )
    add_table(
        doc,
        ["阶段", "主要代码", "作用"],
        [
            ("测点/物理模型", "src/em_core.py", "生成半球面测点、球坐标极化基、等效源网格、偶极子传播响应。"),
            ("CST 数据接口", "src/cst_io.py、src/check_cst_export.py", "读取 nearfield/farfield CSV，检查字段、频点、样本和复数场。"),
            ("远场重建", "src/run_baseline.py、src/run_cst_reconstruction.py", "构建测量矩阵，求等效源系数，再外推远场方向图。"),
            ("Level 1 校准", "src/run_cst_level1_angular_calibration.py", "用 CST FarfieldPlot 数据检查角域方向图一致性。"),
            ("识别", "src/run_cst_recognition.py", "提取空间-频谱-极化特征，训练 SVM/RF 分类器。"),
            ("消融", "src/run_cst_recognition_ablation.py", "减少测点或频点，检查识别是否仍超过 85%。"),
            ("结构对照", "src/run_cst_structure_comparison.py", "施加简化结构遮挡，检查方向图变化和跨域识别精度。"),
        ],
        widths=[1.2, 2.3, 2.75],
        font_size=8.5,
    )

    add_heading(doc, "2.2 半球测点与少测点算法", 2)
    add_para(
        doc,
        "半球测点由 make_hemisphere_layout() 生成。它根据半径、俯仰角和方位角，把测点均匀铺在上半球面上，并同时生成每个测点的球坐标极化基。当前使用 162 个空间点，半径 13 m。这个点数不是密集采样，而是一个工程折中：点数足够覆盖 2π 立体角，又不会让测量系统或 CST 数据量过重。",
    )
    add_para(
        doc,
        "少测点算法主要用 farthest_point_subset()。直白说，它会先选一个点，再不断选择“离已选点最远”的新点，使子集尽量均匀覆盖半球面。这样比随机删点更不容易出现某个方向完全没测到的问题。",
    )
    add_ref(doc, "L07 Valdez/Yuffa/Wakin 受限域压缩感知；L08 Bangun & Culotta-Lopez 球面测点优化；L25 多频稀疏测量。")

    add_heading(doc, "2.3 等效源反演与远场重建算法", 2)
    add_para(
        doc,
        "重建算法的核心思想是“外部测量场可以由内部一组等效源产生”。代码先在被测体包络内部放置等效源网格，每个等效源可以理解成一个待求的辐射小单元。然后计算每个等效源传播到每个测点会产生什么电场，组成测量矩阵 A。",
    )
    add_para(
        doc,
        "如果把半球面测到的复电场记为 y，把未知等效源强度记为 x，那么问题就是 y = A x。实际测量有噪声、测点有限，不能直接硬求逆，所以代码使用 Tikhonov 正则化：求一个既能解释测量数据、又不会过度振荡的 x。得到 x 后，再用这些等效源向远区角域外推，得到远场方向图。",
    )
    add_table(
        doc,
        ["算法步骤", "对应函数", "作用"],
        [
            ("生成等效源网格", "make_equivalent_grid()", "在 12m 级目标包络内部放置候选等效源位置。"),
            ("计算传播响应", "sensor_response()、build_measurement_matrix()", "计算每个等效源到每个测点的复电场贡献。"),
            ("正则化反演", "solve_tikhonov()", "从测量数据反推出等效源系数，避免病态求逆。"),
            ("远场外推", "farfield_pattern()", "由反演得到的等效源计算远区辐射方向图。"),
            ("误差评价", "pattern_metrics()", "输出 NMSE、相关系数、主瓣误差、峰值误差。"),
        ],
        widths=[1.45, 2.15, 2.65],
        font_size=8.6,
    )
    recon = data["recon"]
    rows = []
    for name in ["full_100", "optimized_75", "optimized_50", "random_50", "random_25"]:
        item = next((r for r in recon if r["experiment"] == name), None)
        if item:
            rows.append(
                (
                    name,
                    item["sensor_points"],
                    item["measurement_channels"],
                    fmt_float(item["nmse"], 4),
                    fmt_float(item["correlation"], 6),
                    fmt_float(item["main_lobe_error_deg"], 4) + " deg",
                )
            )
    add_table(
        doc,
        ["实验", "测点", "通道", "NMSE", "相关系数", "主瓣误差"],
        rows,
        widths=[1.25, 0.7, 0.75, 1.05, 1.25, 1.25],
        font_size=8.7,
    )
    add_figure(
        doc,
        ROOT / "outputs" / "baseline" / "farfield_comparison_full_100.png",
        "图 6  全测点条件下的远场重建对比。重建方向图与参考方向图高度一致。",
        width=5.9,
    )
    add_figure(
        doc,
        ROOT / "outputs" / "baseline" / "sampling_tradeoff.png",
        "图 7  测点数量与重建误差的折中。50% 优化测点仍能保持较高相关性，体现少测点验证价值。",
        width=5.9,
    )
    add_para(
        doc,
        "这些结果说明：在当前基线仿真中，162 点全测点可以稳定重建远场；采用几何均匀的 50% 优化测点时，相关系数仍约 0.9986，主瓣误差仍为 0 deg。也就是说，代码已经形成了“先高精度重建，再减少测点验证”的完整算法链。",
    )
    add_ref(doc, "L03 Yaghjian；L05 Kornprobst；L10 Sarkar & Taaghol；L12 Regue 等效偶极子源识别与远场预测。")

    add_heading(doc, "2.4 CST 数据读取与校验算法", 2)
    add_para(
        doc,
        "CST 导出的数据不能直接默认可信，必须先做格式和物理一致性检查。cst_io.py 负责把 nearfield/farfield CSV 读入，检查 sample_id、frequency、sensor_id、坐标、复数字段是否齐全，并把 Ex/Ey/Ez 必要时转换到 theta/phi 极化方向。",
    )
    add_table(
        doc,
        ["检查对象", "检查内容", "为什么需要"],
        [
            ("nearfield", "测点坐标、频点、复电场、样本编号", "保证每个样本在 162 点上都有可用测量数据。"),
            ("farfield", "theta/phi 角度、远场复数或功率", "保证重建结果有可对比的远场参考。"),
            ("nearfield + farfield 配对", "样本和频点是否一一对应", "避免拿 A 样本近场去对 B 样本远场。"),
            ("极化转换", "Ex/Ey/Ez -> theta/phi", "把 CST 笛卡尔分量转换成方向图和识别更常用的球面极化分量。"),
        ],
        widths=[1.35, 2.25, 2.65],
        font_size=8.7,
    )
    add_para(
        doc,
        f"当前 Level 2 数据校验结果显示：样本数 {data['recog']['feature_metadata']['sample_count']}，频点数 {len(data['recog']['feature_metadata']['frequency_hz'])}，测点数 {data['recog']['feature_metadata']['sensor_count']}，特征数 {data['recog']['feature_metadata']['feature_count']}，数据链路已满足识别算法输入要求。",
    )
    add_ref(doc, "L01 IEEE 149 和 L02 IEEE 1720 都强调测量坐标、极化、数据交付格式和校准链路必须明确。")

    add_heading(doc, "2.5 空间-频谱特征识别算法", 2)
    add_para(
        doc,
        "识别算法的输入不是单个方向的幅度，而是全景辐射数据。每个样本都包含多个频点、162 个空间测点和多个场分量。代码把这些数据变成空间-频谱-极化联合特征，用来区分不同源配置和运行状态。",
    )
    add_table(
        doc,
        ["特征块", "直白解释", "作用"],
        [
            ("归一化 log 幅度", "看每个测点、每个频点场强如何分布", "反映辐射空间方向性和频率响应。"),
            ("相对相位 cos/sin", "把相位差用 cos 和 sin 表示", "保留相位结构，同时避免相位跳变。"),
            ("功率/极化摘要", "统计不同分量、不同极化的能量比例", "反映源方向、姿态和耦合差异。"),
            ("空间功率摘要", "计算空间能量质心、分布离散程度等", "把“哪里更亮”变成可分类特征。"),
            ("频率变化特征", "比较 900-1500 MHz 内能量随频率变化", "体现不同设备或状态的频谱指纹。"),
        ],
        widths=[1.45, 2.35, 2.45],
        font_size=8.6,
    )
    add_para(
        doc,
        "分类器使用两类模型：RBF-SVM 和 Random Forest。SVM-RBF 适合中小样本、高维、非线性分类；Random Forest 作为对照模型，用来检查结果是不是只依赖某一个模型。当前最佳模型是 svm_rbf。",
    )
    r = data["recog"]["recognition"]
    add_table(
        doc,
        ["项目", "当前结果"],
        [
            ("样本数", str(data["recog"]["feature_metadata"]["sample_count"])),
            ("类别", "comm_pair、mixed_avionics、multi_state_on、radar_top"),
            ("训练/测试样本", f"{r['train_sample_count']} / {r['test_sample_count']}"),
            ("最佳模型", r["best_model"]),
            ("最佳准确率", f"{r['best_accuracy']:.3f}"),
            ("赛题要求", "不低于 85%"),
        ],
        widths=[1.7, 4.55],
        font_size=8.9,
    )
    add_figure(
        doc,
        ROOT / "outputs" / "cst_recognition_level2" / "cst_recognition_confusion_matrix.png",
        "图 8  Level 2 四类空间-频谱特征识别混淆矩阵。当前 CST-derived 样本测试准确率为 1.000。",
        width=5.45,
    )
    add_para(
        doc,
        "这个结果说明，当前构造的多源多状态样本在空间分布、频率响应和极化相位上有清晰差异，特征工程能够把这些差异提取出来，并支撑超过 85% 的识别指标。",
    )
    add_ref(doc, "L17 RF 指纹综述；L18 雷达发射机机理特征；L19 真实雷达 SEI；L21 复数 RF 指纹识别。")

    add_heading(doc, "2.6 消融实验：减少测点和频点后是否还能识别", 2)
    add_para(
        doc,
        "消融实验的目的很直接：如果只有全 162 点、全 5 频点才能识别，那么方案工程压力较大；如果减少测点或频点后仍能识别，说明空间-频谱特征有冗余，也说明测量系统有优化空间。",
    )
    ab_rows = []
    for item in data["ablation"]["ablation_cases"]:
        ab_rows.append(
            (
                item["case"],
                str(item["sensor_count"]),
                str(item["frequency_count"]),
                str(item["measurement_channels_per_sample"]),
                f"{item['best_accuracy']:.3f}",
            )
        )
    add_table(
        doc,
        ["消融场景", "测点", "频点", "每样本通道", "准确率"],
        ab_rows,
        widths=[2.0, 0.75, 0.75, 1.25, 1.0],
        font_size=8.1,
    )
    add_figure(
        doc,
        ROOT / "outputs" / "cst_recognition_level2_ablation" / "recognition_ablation_accuracy.png",
        "图 9  Level 2 测点/频点消融结果。当前 8 个消融场景准确率均为 1.000。",
        width=5.9,
    )
    add_para(
        doc,
        "这说明当前样本的类别差异比较稳定：即使把测点降到 50%、频点降到 1 个中心频点，识别结果仍超过 85%。正式汇报时要注意，这并不代表真实外场任何情况下都只需这么少的点，而是说明在当前 CST-derived 数据上，少测点和少频点具有可行性。",
    )
    add_ref(doc, "L07、L08 和 L25 共同支撑少测点/多频稀疏测量；L17-L19 支撑多特征 RF 指纹识别。")

    add_heading(doc, "2.7 结构扰动算法与结果含义", 2)
    add_para(
        doc,
        "结构扰动代码用几何遮挡模型近似机体安装影响。它用 ray-box 和 ray-ellipsoid 计算从源到不同方向的路径是否穿过机身、机翼或尾翼，再把遮挡量转换为方向相关的 dB 衰减，施加到 nearfield/farfield 数据上。",
    )
    add_table(
        doc,
        ["函数", "作用"],
        [
            ("ray_box_length()、ray_ellipsoid_length()", "计算射线穿过简化机身/机翼/尾翼的路径长度。"),
            ("source_shadow_db()", "把遮挡路径长度转换为方向相关的衰减。"),
            ("apply_structure_to_farfield()", "生成结构扰动后的远场表。"),
            ("apply_structure_to_nearfield()", "生成结构扰动后的近场表。"),
            ("recognition_domain_shift()", "用无结构数据训练、结构扰动数据测试，检查识别鲁棒性。"),
        ],
        widths=[2.25, 4.0],
        font_size=8.7,
    )
    add_para(
        doc,
        "结果显示结构扰动确实会改变方向图，平均遮挡约 3.06 dB，平均方向图相关系数约 0.730；但在当前四类样本上，跨域识别准确率仍为 1.000。这可以作为“结构会影响辐射空间分布，但当前特征识别链路仍稳定”的证据。",
    )
    add_ref(doc, "L05、L06 支撑复杂结构和等效表示之间需要边界验证；该部分是有界结构敏感性分析，不是 full-wave 机体散射仿真。")

    add_heading(doc, "2.8 当前可以汇报的结论与边界", 2)
    add_table(
        doc,
        ["可以直接讲", "需要保守讲"],
        [
            (
                "已经形成 CST 工程、CSV 数据、Python 重建、识别、消融和结构对照的闭环。",
                "当前 Level 2 是 CST-derived element-library 叠加样本，不是完整复杂机体 full-wave 解。",
            ),
            (
                "13 m 半球 162 点布局支撑 2π 空域采样，并已写入 Level 1 CST 探针工程。",
                "162 点是稀疏工程布局，不是波长级高密度近场扫描。",
            ),
            (
                "基线重建 full_100 NMSE 约 2.97e-4，相关系数约 0.9995；optimized_50 仍约 0.9986。",
                "重建高精度结论主要来自基线仿真和 Level 1 角域校准，复杂机体 full-wave 近场反演仍可后续加强。",
            ),
            (
                "Level 2 四类识别准确率 1.000，超过 85% 要求；消融后仍稳定。",
                "真实外场数据、更多载体类别、噪声和姿态变化需要后续实测或更复杂仿真继续验证。",
            ),
        ],
        widths=[3.1, 3.15],
        font_size=8.6,
    )


def add_appendix(doc: Document):
    add_heading(doc, "附录：本报告引用的关键文件", 1)
    add_table(
        doc,
        ["用途", "本机位置"],
        [
            ("Level 1 CST 工程", rel(ROOT / "outputs" / "cst_solver_ready_level1_projects" / "projects")),
            ("Level 2 CST 基元库", rel(ROOT / "outputs" / "cst_level2_element_library" / "projects")),
            ("CST Level 1 数据", rel(ROOT / "data" / "cst_exports" / "level1")),
            ("CST Level 2 数据", rel(ROOT / "data" / "cst_exports" / "level2")),
            ("重建算法核心", rel(ROOT / "code" / "em_core.py")),
            ("CST 数据接口", rel(ROOT / "code" / "cst_io.py")),
            ("识别算法", rel(ROOT / "code" / "run_cst_recognition.py")),
            ("消融算法", rel(ROOT / "code" / "run_cst_recognition_ablation.py")),
            ("结构对照算法", rel(ROOT / "code" / "run_cst_structure_comparison.py")),
            ("文献矩阵", rel(ROOT / "docs" / "literature_matrix.md")),
        ],
        widths=[2.0, 4.25],
        font_size=8.4,
    )
    add_para(
        doc,
        "文献编号沿用 docs/literature_matrix.md：L01 IEEE 149，L02 IEEE 1720，L03 Yaghjian，L05 Kornprobst，L06 Malmstrom，L07 Valdez/Yuffa/Wakin，L08 Bangun & Culotta-Lopez，L17-L19 RF/雷达指纹识别，L25 多频稀疏测量。",
        size=9.2,
        color=MUTED,
    )


def collect_data() -> dict:
    return {
        "l1_merge": read_json(ROOT / "outputs" / "cst_level1_merge_report" / "level1_merge_summary.json"),
        "l1_ang": read_json(ROOT / "outputs" / "cst_level1_angular_calibration" / "angular_calibration_summary.json"),
        "structure": read_json(ROOT / "outputs" / "cst_structure_comparison" / "structure_comparison_summary.json"),
        "recon": read_csv_rows(ROOT / "outputs" / "baseline" / "reconstruction_metrics.csv"),
        "recog": read_json(ROOT / "outputs" / "cst_recognition_level2" / "cst_recognition_metrics.json"),
        "ablation": read_json(ROOT / "outputs" / "cst_recognition_level2_ablation" / "recognition_ablation_summary.json"),
    }


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = collect_data()
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_cst_section(doc, data)
    add_code_section(doc, data)
    add_appendix(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
