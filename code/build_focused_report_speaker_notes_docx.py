# -*- coding: utf-8 -*-
from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "initial_report_focused"
DOCX_PATH = OUT_DIR / "复杂航空载体电磁辐射空域特性测量_初期汇报_聚焦版_讲稿.docx"

FONT_CN = "Microsoft YaHei"
FONT_EN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(28, 28, 28)
MUTED = RGBColor(96, 96, 96)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_YELLOW = "FFF8E6"
PALE_GREEN = "EEF7EE"


def read_json(path: str | Path) -> dict:
    with Path(path).open("r", encoding="utf-8") as f:
        return json.load(f)


def read_csv_rows(path: str | Path) -> list[dict[str, str]]:
    with Path(path).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def rel(path: str | Path) -> str:
    try:
        return str(Path(path).resolve().relative_to(ROOT)).replace("/", "\\")
    except Exception:
        return str(path).replace("/", "\\")


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def set_paragraph_font(paragraph, size=10.5, color=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def write_cell(cell, text: str, size=9.0, bold=False, color=None, fill: str | None = None):
    cell.text = ""
    if fill:
        set_cell_shading(cell, fill)
    set_cell_margins(cell)
    lines = str(text).split("\n")
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color or INK)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[str]], widths: list[float], font_size=8.8):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        write_cell(cell, header, size=font_size, bold=True, fill=LIGHT_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            write_cell(cells[idx], str(value), size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_para(doc: Document, text: str = "", size=10.5, bold=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color or INK)
    return p


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    return p


def add_note_box(doc: Document, title: str, body: str, fill: str = PALE_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=150, bottom=130, end=150)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    for line in body.split("\n"):
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.18
        r2 = p2.add_run(line)
        set_run_font(r2, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_speech_block(doc: Document, location: str, topic: str, script: str, how_to_read: str, terms: str = ""):
    add_heading(doc, topic, 3)
    rows = [
        ("报告位置", location),
        ("可以照着讲", script),
        ("图表怎么看", how_to_read),
    ]
    if terms:
        rows.append(("术语解释", terms))
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for ridx, (label, detail) in enumerate(rows):
        label_cell = table.rows[ridx].cells[0]
        detail_cell = table.rows[ridx].cells[1]
        label_cell.width = Inches(1.18)
        detail_cell.width = Inches(5.32)
        write_cell(label_cell, label, size=8.8, bold=True, fill=LIGHT_BLUE)
        write_cell(detail_cell, detail, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = FONT_EN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    header = section.header.paragraphs[0]
    header.text = "复杂航空载体电磁辐射空域特性测量技术比赛 | 初期汇报聚焦版讲稿"
    set_paragraph_font(header, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "对照报告：复杂航空载体电磁辐射空域特性测量_初期汇报_聚焦版.docx"
    set_paragraph_font(footer, size=8, color=MUTED)


def fmt_float(value, digits=4):
    return f"{float(value):.{digits}g}"


def collect_data() -> dict:
    return {
        "l1_merge": read_json(ROOT / "outputs" / "cst_level1_merge_report" / "level1_merge_summary.json"),
        "l1_ang": read_json(ROOT / "outputs" / "cst_level1_angular_calibration" / "angular_calibration_summary.json"),
        "recon": read_csv_rows(ROOT / "outputs" / "baseline" / "reconstruction_metrics.csv"),
        "recog": read_json(ROOT / "outputs" / "cst_recognition_level2" / "cst_recognition_metrics.json"),
        "ablation": read_json(ROOT / "outputs" / "cst_recognition_level2_ablation" / "recognition_ablation_summary.json"),
        "structure": read_json(ROOT / "outputs" / "cst_structure_comparison" / "structure_comparison_summary.json"),
    }


def recon_value(rows, name, key):
    for row in rows:
        if row["experiment"] == name:
            return row[key]
    return ""


def add_cover(doc: Document):
    add_para(doc, "初期汇报聚焦版讲稿", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "逐页对照报告内容、图表读法与专业名词解释", size=14.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    add_table(
        doc,
        ["项目", "说明"],
        [
            ("适用材料", "《复杂航空载体电磁辐射空域特性测量_初期汇报_聚焦版.docx》"),
            ("讲稿用途", "可直接用于初期汇报发言，也可作为 PPT/答辩前的背稿材料。"),
            ("建议时长", "12-18 分钟；如果时间较短，可只讲每节“可以照着讲”的第一段和每张图的结论句。"),
            ("讲述重点", "先说明 CST 建模产物和用途，再说明代码算法如何用这些数据完成重建、识别和验证。"),
            ("生成时间", "2026-05-29"),
        ],
        widths=[1.5, 5.0],
        font_size=9.2,
    )
    add_note_box(
        doc,
        "汇报总口径",
        "这项工作目前不是单纯画了几张 CST 图，而是形成了“CST 建模生成电磁场数据 -> Python 数据校验 -> 远场重建 -> 空间-频谱识别 -> 消融和结构扰动验证”的初期工程闭环。\n"
        "汇报时要主动说明边界：Level 2 是 CST-derived element-library 叠加样本，不是完整复杂机体 full-wave 全波仿真。这样讲更稳，也更专业。",
        fill=PALE_GREEN,
    )


def add_quick_overview(doc: Document):
    add_heading(doc, "一、开场总述", 1)
    add_note_box(
        doc,
        "30 秒版本",
        "各位老师好，我这次汇报主要讲两部分：第一是 CST 建模，我们如何用标准源、三方向基元库和结构扰动对照来生成可信的电磁场数据；第二是代码算法，我们如何把这些数据用于远场重建、少测点验证和空间-频谱特征识别。当前项目已经打通初期闭环，识别结果超过赛题 85% 指标，同时我们也保守说明了 full-wave 复杂机体仿真仍是后续增强项。",
        fill=PALE_GREEN,
    )
    add_table(
        doc,
        ["汇报阶段", "一句话讲法", "时间建议"],
        [
            ("开场", "说明本次只讲 CST 建模与代码算法，不单独展开文献矩阵。", "1 分钟"),
            ("CST 建模", "讲清楚 Level 1、Level 2、结构对照分别建了什么、为什么这样建、得到什么数据。", "5-7 分钟"),
            ("代码算法", "讲清楚半球测点、等效源反演、远场重建、特征识别和消融实验。", "6-8 分钟"),
            ("边界与后续", "主动说明当前证据边界和下一步 full-wave/实测增强方向。", "1-2 分钟"),
        ],
        widths=[1.35, 4.2, 0.95],
        font_size=8.8,
    )


def add_page_by_page_script(doc: Document, data: dict):
    add_heading(doc, "二、逐页对照讲稿", 1)

    add_speech_block(
        doc,
        "报告第 1 页：封面、阅读导引、一、CST 建模部分开头",
        "0. 封面与本次汇报范围",
        "这一版初期汇报我重新聚焦成两条主线：CST 建模和代码算法。CST 部分回答“我们建了哪些电磁模型、数据从哪里来”；代码部分回答“这些数据怎样进入算法，怎样完成重建和识别”。文献不再单独占一章，而是在每个技术点后标注对应依据，这样更适合初期汇报时快速讲清工程进展。",
        "第 1 页不要逐字读表。重点看“阅读导引”框：它告诉听众这不是完整最终报告，而是一份帮助理解工程链路的讲解稿。讲完这一页后，听众应该知道后面只围绕 CST 和算法展开。",
        "CST：电磁仿真软件，用来建立天线/辐射源模型并求解电磁场。\n工程闭环：不是只做一个环节，而是从建模、导出、算法处理到结果验证都能跑通。",
    )

    add_speech_block(
        doc,
        "报告第 1-2 页：1.1 为什么 CST 建模要分三层，三层建模表",
        "1. CST 建模为什么采用三层结构",
        "我们没有一开始就直接做完整复杂机体 full-wave 仿真，而是采用分层方式。第一层 Level 1 是标准源，用短偶极子和半波偶极子校准坐标、极化和导出链路；第二层 Level 2 是三方向短偶极子基元库，用于合成多源、多频、多状态识别样本；第三层是结构对照，用简化机身、机翼、尾翼遮挡模型评估结构安装对方向图和识别的影响。\n\n这样做的原因是：复杂机体 full-wave 求解成本高，而且一旦出错很难判断是模型问题、坐标问题、极化问题还是数据导出问题。先用标准源把基础链路校准，再扩展到多源识别，最后再看结构扰动，这条路线更容易复现和解释。",
        "看三层建模表时，横向看四列：层级、CST 建模内容、主要回答的问题、当前产物。讲表时不要每个字都读，只抓住三个关键词：Level 1 校准链路，Level 2 构造识别数据，结构对照说明边界。",
        "full-wave：全波仿真，指直接求解麦克斯韦方程的电磁仿真，精度高但计算量大。\nelement-library：基元库，先仿真几个基本辐射单元，再组合成复杂样本。\n标准源：物理规律明确、适合校准的简单辐射源，例如偶极子。",
    )

    add_speech_block(
        doc,
        "报告第 2-3 页：1.2 测量布局如何体现在 CST 中，图 1",
        "2. 半球面测量布局和图 1 的讲法",
        "测量方案的源头在 Python，代码统一生成 13 m 半径、覆盖上半空间 2π 立体角的 162 个半球面测点。每个点有固定编号和坐标，因此同一套测点可以被 CST 工程、CSV 数据和 Python 算法共同使用。Level 1 中这 162 个点已经作为 efarfield 探针写入 CST 工程；Level 2 为了控制求解规模，采用基元库加后处理叠加，但输出数据仍然按同一套 162 点布局组织。\n\n这里要强调：162 点不是波长级高密度近场扫描，而是面向等效源重建和特征识别的稀疏工程布局。它的价值在于覆盖 2π 空间，同时控制布设和数据规模。",
        "图 1 是三维半球测点图。看图时先看半球形状，说明覆盖的是被测体上方 2π 空间；再看彩色点，说明这些就是空间测量点；最后看中间红色框线，说明它对应 12m x 10m x 8m 级别被测包络。图 1 的结论句是：我们的测量布局能够包络赛题规定尺寸，并提供多角度空间采样。",
        "2π 立体角：完整球面是 4π，上半球就是 2π。\n近场测点：在目标周围某些空间位置采集电场的位置。\nefarfield 探针：CST 中用于提取场或远场相关响应的探针设置。",
    )

    l1_merge = data["l1_merge"]
    l1_ang = data["l1_ang"]
    add_speech_block(
        doc,
        "报告第 3-5 页：1.3 Level 1 标准源，图 2、图 3、Level 1 结果表",
        "3. Level 1 标准源建模、结果和用途",
        f"Level 1 的目标是先校准链路。我们建立了两个 required 标准源：1.2 GHz z 向短偶极子和 z 向半波偶极子。选择偶极子的原因是它的方向图规律明确，容易检查坐标、极化和主瓣方向是否正确。当前 required cases 完成 {l1_merge['required_complete_cases']}/{l1_merge['required_cases']}，合并 nearfield 行数为 {l1_merge['merged_nearfield_rows']}，farfield 行数为 {l1_merge['merged_farfield_rows']}。\n\nLevel 1 的角域校准结果比较好：max NMSE 约 {fmt_float(l1_ang['max_nmse'], 4)}，min correlation 约 {fmt_float(l1_ang['min_correlation'], 6)}，主瓣误差为 {l1_ang['max_main_lobe_error_deg']} deg。这说明 CST 的标准源远场结果和我们后处理中的角域模型高度一致，坐标、极化和导出链路没有明显错误。",
        "图 2 是 CST 中打开的半波偶极子工程截图。看图时只需要说明：这证明工程文件是真实可打开的 CST 产物，而不是只存在于代码里的假想模型。\n\n图 3 是角域校准图。通常左边是 CST 或参考结果，中间是校准/拟合结果，右边是差异图。看图时重点看右边差异是否接近均匀、数值是否小；结合表格里的 NMSE 和相关系数，说明二者重合度很高。",
        "NMSE：归一化均方误差，越小越好。\ncorrelation：相关系数，越接近 1 越说明两张方向图形状越像。\n主瓣误差：最大辐射方向偏差，越接近 0 越好。\nnearfield/farfield：近场是靠近目标处的测量场，远场是远距离角域辐射分布。",
    )

    add_speech_block(
        doc,
        "报告第 5-6 页：1.4 Level 2 基元库，图 4，Level 2 建模设置表",
        "4. Level 2 基元库和多源多状态样本",
        "Level 2 的目标是构造识别算法需要的数据集。我们没有把完整复杂载体和所有辐射源一次性放进 CST 里求解，而是先建立 x、y、z 三方向短偶极子基元库，每个基元覆盖 900、1050、1200、1350、1500 MHz 五个频点。之后在 Python 中根据源位置、方向、幅度和相位，把这些基元叠加成 48 个多源多状态样本。\n\n这条路线的优点是算力压力小、可解释、可复现。每个样本都能追溯到哪些等效辐射基元参与了叠加，后续识别结果如果异常，也能回头检查源组合和频点。",
        "图 4 是 Level 2 x 向短偶极子基元在 CST 中打开的截图。讲图时要强调它不是最终复杂飞机模型，而是“基元库中的一个基本构件”。表格中重点看三行：三方向短偶极子、五个频点、48 个样本。结论句是：Level 2 用可控基元库生成了多频、多源、多状态识别数据。",
        "多源多状态：不同辐射源组合、不同工作状态形成不同样本。\nsample-frequency：一个样本在一个频点上的一组数据；48 个样本乘 5 个频点就是 240 组。\n叠加：利用电磁场线性近似，把多个基元场按幅度相位组合起来。",
    )

    s = data["structure"]
    add_speech_block(
        doc,
        "报告第 6-7 页：1.5 结构对照，图 5，结构指标表",
        "5. 结构遮挡对照和边界说明",
        f"为了避免把 Level 2 的基元叠加结果误说成完整机体 full-wave 仿真，我们增加了结构遮挡对照。这个对照不是完整机体 CST 全波求解，而是在 Level 2 场数据基础上加入可复现的机身、机翼、尾翼遮挡转移函数。结果显示平均遮挡约 {s['mean_shadow_db']:.2f} dB，P95 遮挡约 {s['p95_shadow_db']:.2f} dB，最大遮挡约 {s['max_shadow_db']:.2f} dB，平均方向图相关系数约 {s['mean_pattern_correlation']:.3f}，跨域识别准确率仍为 {s['cross_domain_accuracy']:.3f}。\n\n这部分最重要的作用是说明边界：结构确实会改变方向图，但在当前样本上识别链路仍然稳定。后续如果时间和算力允许，可以再做更完整的复杂机体 full-wave 验证。",
        "图 5 一般看三块：无遮挡方向图、加入结构后的方向图、二者差异。看图时先说明颜色代表辐射强弱，再看差异图中哪些角度变化明显。表格里重点看 shadow dB 和 cross-domain accuracy。结论句是：结构影响不可忽略，但当前识别算法在这个扰动下仍保持稳定。",
        "shadow dB：遮挡导致的场强衰减，dB 越大说明影响越强。\nP95：95% 情况下不超过这个数，常用来描述大多数情况的上界。\ncross-domain accuracy：在一种数据域训练，在另一种扰动数据域测试的准确率，用来衡量鲁棒性。",
    )

    add_speech_block(
        doc,
        "报告第 7-8 页：二、代码算法部分，2.1 代码总流程",
        "6. 代码算法总流程",
        "代码部分可以按数据流讲：先生成半球测点和等效源网格，再读取 CST 或仿真得到的近场/远场数据；随后一条支路做等效源反演和远场重建，另一条支路做空间-频谱特征识别；最后通过消融实验和结构扰动对照验证少测点、少频点和结构变化下的稳定性。\n\n这里要让听众明白：代码不是只训练了一个分类器，而是承担了测点生成、物理反演、CST 数据校验、识别建模和结果审计五类工作。",
        "第 7-8 页主要看代码职责表。讲表时按阶段讲，不要逐行读路径。可以说：em_core.py 是物理核心，cst_io.py 是数据接口，run_cst_recognition.py 是识别，ablation 和 structure comparison 是验证算法稳定性的脚本。",
        "数据流：数据从哪里来、经过哪些处理、最后输出什么结果。\n审计：检查数据、结果和文件是否一致，避免只看单个漂亮指标。",
    )

    add_speech_block(
        doc,
        "报告第 8-10 页：2.2-2.3，重建算法表，图 6、图 7",
        "7. 等效源反演与远场重建",
        f"远场重建的核心思想是等效源。我们不直接猜复杂载体内部每个真实辐射细节，而是在被测体包络内放置一组等效源。外部半球面测到的复电场 y，可以看成传播矩阵 A 乘以等效源系数 x，也就是 y = A x。代码通过 Tikhonov 正则化反演求出 x，然后再由这些等效源向远区外推得到方向图。\n\n当前基线结果中，full_100 使用 162 个测点，NMSE 约 {fmt_float(recon_value(data['recon'], 'full_100', 'nmse'), 4)}，相关系数约 {fmt_float(recon_value(data['recon'], 'full_100', 'correlation'), 6)}。optimized_50 使用 81 个测点，相关系数仍约 {fmt_float(recon_value(data['recon'], 'optimized_50', 'correlation'), 6)}，主瓣误差为 {fmt_float(recon_value(data['recon'], 'optimized_50', 'main_lobe_error_deg'), 3)} deg。这说明在当前仿真条件下，减少测点后仍能保持较好远场推算能力。",
        "图 6 看三块：左边参考方向图，中间重建方向图，右边差异图。左中越像、右边越接近浅色或均匀，就说明重建越好。\n\n图 7 看横轴不同测点方案，纵轴看误差或相关性变化。重点不是说点越少永远越好，而是说明 optimized_50 比随机少点更稳，体现测点选择算法有价值。",
        "等效源：用一组数学上的源代替真实复杂辐射源。\n传播矩阵 A：描述每个等效源对每个测点贡献多少电场。\nTikhonov 正则化：反演时加入稳定项，避免因为测点少或噪声导致结果乱跳。\n复电场：同时包含幅度和相位的电场数据。",
    )

    add_speech_block(
        doc,
        "报告第 10-11 页：2.4 CST 数据读取与校验，2.5 识别特征表",
        "8. CST 数据校验与空间-频谱特征",
        "CST 导出的数据不能直接拿来训练或重建，必须先检查。cst_io.py 会检查 nearfield 和 farfield 的字段是否完整，样本和频点是否匹配，坐标和复数电场是否可用。如果 CST 输出的是 Ex/Ey/Ez 笛卡尔分量，代码还可以转换到 theta/phi 球面极化分量。\n\n识别特征不是单个方向的场强，而是把“空间、频率、极化、相位”一起建模。空间特征看哪个方向辐射更强，频谱特征看不同频点能量如何变化，极化特征看不同场分量比例，相位特征保留传播差异。这样形成的是空间-频谱-极化联合指纹。",
        "第 10-11 页的表格主要看三件事：数据是否完整、特征有哪些、每类特征解释什么物理差异。讲特征表时可以用“哪里亮、随频率怎么变、极化比例怎样、相位差怎样”四句话解释，听众会更容易理解。",
        "Ex/Ey/Ez：笛卡尔坐标下三个方向的电场分量。\ntheta/phi：球坐标下两个切向极化方向，更适合描述球面方向图。\n空间-频谱-极化指纹：把空间分布、频率响应和极化相位合起来作为识别特征。",
    )

    recog = data["recog"]
    add_speech_block(
        doc,
        "报告第 11-12 页：2.5 识别结果，图 8",
        "9. SVM/RF 识别算法和混淆矩阵",
        f"识别算法使用了两类模型：RBF-SVM 和 Random Forest。SVM-RBF 适合中小样本、高维、非线性分类；Random Forest 作为对照，检查结果是否依赖单一模型。当前 Level 2 数据有 {recog['feature_metadata']['sample_count']} 个样本、5 个频点、162 个测点，构成 {recog['feature_metadata']['feature_count']} 维特征。最佳模型是 {recog['recognition']['best_model']}，测试准确率为 {recog['recognition']['best_accuracy']:.3f}，超过赛题要求的 85%。\n\n这说明当前四类源配置或运行状态在空间-频谱特征上差异明显，算法能够把它们区分开。",
        "图 8 是混淆矩阵。看法很简单：纵轴是真实类别，横轴是预测类别；颜色集中在对角线，说明预测正确；如果非对角线上有数字，就代表某一类被错分成另一类。当前图中主要数字都在对角线上，所以四类都能区分。",
        "SVM-RBF：一种支持向量机分类器，RBF 核可以处理非线性边界。\nRandom Forest：随机森林，多个决策树投票分类。\n混淆矩阵：显示每个真实类别被预测成哪些类别的表。",
    )

    add_speech_block(
        doc,
        "报告第 12-13 页：2.6 消融实验，图 9",
        "10. 消融实验和图 9 的讲法",
        "消融实验回答一个问题：如果减少测点或减少频点，识别还能不能保持？这对工程很重要，因为真实测量中测点越少，布设和采集成本越低。当前做了 8 个场景，包括全 5 频点、75% 测点、50% 测点、25% 测点，以及只用 3 个频点或 1 个中心频点的情况。当前这些消融场景的准确率均为 1.000。\n\n这说明当前仿真样本中，类别差异比较稳定，少测点和少频点仍然保留了足够识别信息。不过汇报时要保守说明：这不代表真实外场任何情况下都能这么少，而是说明当前 CST-derived 数据上具有可行性。",
        "图 9 是柱状图。横轴是不同消融场景，纵轴是准确率，红色虚线是赛题 85% 阈值。所有蓝色柱都高于红线，并且达到 1.000，这就是这张图的核心结论。",
        "消融实验：故意去掉一部分输入条件，看结果是否下降。\n85% 阈值：赛题要求的空间频率特征辨识精度下限。\n鲁棒性：条件变化后结果是否仍稳定。",
    )

    add_speech_block(
        doc,
        "报告第 13-14 页：2.7 结构扰动算法，2.8 可以汇报的结论与边界",
        "11. 结构扰动算法与最终边界",
        "结构扰动算法用简化几何来模拟机体安装影响。代码会计算从辐射源到某个方向的路径是否穿过机身、机翼或尾翼，再把穿过长度转成方向相关的 dB 衰减，施加到近场和远场数据上。这个结果用于说明：结构确实会改变空间方向图，因此后续完整复杂机体 full-wave 或实测数据仍然有增强价值。\n\n最后第 2.8 节非常关键。汇报结论要分两边讲：可以直接讲的是我们已经形成 CST 工程、CSV 数据、Python 重建、识别、消融和结构对照闭环；需要保守讲的是 Level 2 目前是 CST-derived element-library 叠加样本，不是完整复杂机体 full-wave 解。这样的表达既体现工作量，也避免夸大。",
        "第 14 页的两列表格要重点讲。左列“可以直接讲”是汇报亮点，右列“需要保守讲”是边界。答辩时老师如果追问 full-wave、162 点是否够、准确率 1.000 是否过高，都可以回到这张表回答。",
        "边界：当前结果能证明什么、不能证明什么。\n结构扰动：在已有数据上加入结构遮挡影响，测试算法是否稳定。\nCST-derived：来自 CST 基元或 CST 导出结果再加工的数据，不等同于完整一体化全波仿真。",
    )


def add_glossary(doc: Document):
    add_heading(doc, "三、专业名词速查", 1)
    add_table(
        doc,
        ["名词", "直白解释", "汇报时怎么说"],
        [
            ("2π 空间", "半个球面的空间覆盖。完整球面是 4π，上半球是 2π。", "我们的测点覆盖被测体上方半球空间。"),
            ("半球面测点", "分布在一个半球面上的采样位置。", "162 个点提供多角度观测，不是单方向测量。"),
            ("双极化/多分量", "不同方向上的电场分量。", "保留极化差异能增强识别特征。"),
            ("近场", "靠近被测体处的电磁场。", "传感器采集的是近场或准近场数据。"),
            ("远场", "远距离上只随角度变化的辐射方向图。", "重建目标是推算远场辐射分布。"),
            ("等效源", "用数学源代替真实复杂辐射源。", "它让复杂载体辐射问题变成可反演问题。"),
            ("Tikhonov 正则化", "反演时加入稳定约束。", "用于避免测点有限和噪声导致的求解不稳定。"),
            ("NMSE", "归一化均方误差，越小越好。", "用来衡量重建图和参考图差多少。"),
            ("相关系数", "两个方向图形状相似度，越接近 1 越好。", "0.999 级别说明形状高度一致。"),
            ("主瓣误差", "最大辐射方向的角度偏差。", "0 deg 说明主辐射方向没有偏移。"),
            ("混淆矩阵", "真实类别和预测类别的对应表。", "数字集中在对角线代表分类正确。"),
            ("消融实验", "减少测点/频点后再测试。", "用来证明方案是否有冗余和鲁棒性。"),
            ("SVM-RBF", "一种适合高维非线性分类的机器学习模型。", "用于识别不同源配置和运行状态。"),
            ("Random Forest", "多个决策树投票的分类模型。", "作为 SVM 的对照模型。"),
            ("full-wave", "完整求解麦克斯韦方程的电磁仿真。", "当前复杂机体 full-wave 是后续增强项。"),
        ],
        widths=[1.45, 2.7, 2.35],
        font_size=8.3,
    )


def add_figure_cheatsheet(doc: Document):
    add_heading(doc, "四、图表读法速查", 1)
    add_table(
        doc,
        ["图表", "先看哪里", "一句话结论"],
        [
            ("图 1 半球测点", "看半球形测点和中间目标包络。", "13 m 半球 162 点实现 2π 多角度覆盖。"),
            ("图 2 CST 半波偶极子", "看 CST 界面中的偶极子几何。", "Level 1 是真实可打开的 CST 工程。"),
            ("图 3 角域校准", "看参考图、校准图、差异图是否接近。", "标准源方向图一致性好，链路可信。"),
            ("图 4 Level 2 基元", "看三方向基元库中的 x 向短偶极子。", "基元库用于合成多源多状态样本。"),
            ("图 5 结构对照", "看无遮挡、结构扰动、差异图。", "结构会改变方向图，但识别仍稳定。"),
            ("图 6 远场重建", "看参考图、重建图、差异图。", "全测点重建和参考方向图高度一致。"),
            ("图 7 采样折中", "看不同测点方案下误差和相关性。", "优化少测点比随机少点更稳。"),
            ("图 8 混淆矩阵", "看数字是否集中在对角线。", "四类样本识别准确率达到 1.000。"),
            ("图 9 消融柱状图", "看柱子是否高于 0.85 红线。", "减少测点/频点后仍超过赛题指标。"),
        ],
        widths=[1.55, 2.45, 2.5],
        font_size=8.5,
    )


def add_qa(doc: Document):
    add_heading(doc, "五、可能被问到的问题", 1)
    add_table(
        doc,
        ["问题", "建议回答"],
        [
            (
                "162 个测点算多吗？",
                "不算高密度，属于稀疏工程布局。它的目的不是波长级扫描，而是在 2π 空间覆盖下支撑等效源重建和特征识别，并通过消融实验验证减少测点后的稳定性。",
            ),
            (
                "测量点是在 CST 里布置的吗？",
                "测点坐标由 Python 统一生成。Level 1 中 162 个点已经作为 efarfield 探针写入 CST 工程；Level 2 为控制网格规模，采用基元库加后处理叠加，但数据仍按同一套测点布局组织。",
            ),
            (
                "Level 2 是完整复杂飞机 full-wave 仿真吗？",
                "不是。Level 2 是 CST-derived element-library 叠加样本，用于证明多源多状态识别链路。完整复杂机体 full-wave 是后续增强项。当前报告已单独用结构遮挡对照说明边界。",
            ),
            (
                "准确率 1.000 会不会太理想？",
                "这是当前 CST-derived 仿真样本上的结果，说明样本类别差异明显、特征链路有效。汇报时要同时说明真实外场噪声、姿态变化和更多载体类别仍需后续实测或更复杂仿真验证。",
            ),
            (
                "为什么要做 Level 1 标准源？",
                "因为标准源方向图规律明确，可以先检查坐标、极化、导出和读取是否正确。它是后续复杂样本识别前的校准环节。",
            ),
        ],
        widths=[1.9, 4.6],
        font_size=8.5,
    )


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = collect_data()
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_quick_overview(doc)
    add_page_by_page_script(doc, data)
    add_glossary(doc)
    add_figure_cheatsheet(doc)
    add_qa(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
