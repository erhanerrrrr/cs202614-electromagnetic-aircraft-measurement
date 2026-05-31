from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_MD = ROOT / "docs" / "solution_report_draft.md"
OUT_DIR = ROOT / "submission" / "01_report"
OUT_DOCX = OUT_DIR / "solution_report.docx"
SUMMARY_JSON = ROOT / "outputs" / "final_report" / "final_report_summary.json"

BODY_FONT = "Microsoft YaHei"
LATIN_FONT = "Calibri"
CODE_FONT = "Consolas"
BLUE = RGBColor(31, 77, 120)
ACCENT = RGBColor(46, 116, 181)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"

FIGURES_BY_HEADING = {
    "4.1 半球面布局": [
        (
            "图 1  13 m 半球面 2π 测点布局与 12m x 10m x 8m 包络",
            ROOT / "outputs" / "baseline" / "sensor_layout_hemisphere.png",
        )
    ],
    "5.2 Level 1 required 标准源重建结果": [
        (
            "图 2  Level 1 FarfieldPlot-derived 角域校准对比",
            ROOT
            / "outputs"
            / "cst_level1_angular_calibration"
            / "L1_short_dipole_z_1p2G"
            / "angular_farfield_compare.png",
        )
    ],
    "6.2 当前测点压缩参考结果": [
        (
            "图 3  30 dB 条件下测点数量与重建精度关系",
            ROOT / "outputs" / "reconstruction_robustness" / "reconstruction_sensor_tradeoff_30dB.png",
        )
    ],
    "7.3 Level 2 full48 识别结果": [
        (
            "图 4  Level 2 full48 空-频-极化识别混淆矩阵",
            ROOT / "outputs" / "cst_recognition_level2" / "cst_recognition_confusion_matrix.png",
        )
    ],
    "7.4 识别测点/频点删减验证": [
        (
            "图 5  Level 2 full48 测点/频点删减识别准确率",
            ROOT / "outputs" / "cst_recognition_level2_ablation" / "recognition_ablation_accuracy.png",
        )
    ],
    "7.5 Level 2 简化结构遮挡对照": [
        (
            "图 6  Level 2 简化载体遮挡前后方向图对比",
            ROOT
            / "outputs"
            / "cst_structure_comparison"
            / "plots"
            / "L2_comm_pair_000_1200MHz_structure_compare.png",
        )
    ],
    "9. 误差、鲁棒性与工程可行性分析": [
        (
            "图 7  正则化参数扫描与鲁棒性参考",
            ROOT / "outputs" / "reconstruction_robustness" / "reconstruction_lambda_scan_optimized50_30dB.png",
        )
    ],
}


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
    font: str = BODY_FONT,
) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT if font == BODY_FONT else font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT if font == BODY_FONT else font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 6, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("**", "")
    return text


INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")


def add_inline_runs(paragraph, text: str, *, size: float = 10.6) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token.strip("`"))
            set_run_font(run, size=size, font=CODE_FONT, color=BLUE)
        else:
            run = paragraph.add_run(token.strip("*"))
            set_run_font(run, size=size, bold=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size)


def add_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=10, line=1.0)
    run = p.add_run(caption)
    set_run_font(run, size=9.2, italic=True, color=MUTED)


def add_figure(doc: Document, caption: str, path: Path) -> bool:
    if not path.exists():
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=4, after=8)
        run = p.add_run(f"[图表源文件缺失：{path.relative_to(ROOT)}]")
        set_run_font(run, size=9.5, italic=True, color=RGBColor(155, 28, 28))
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=8, after=2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.9))
    add_caption(doc, caption)
    return True


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.35])
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=2, after=2, line=1.15)
    add_inline_runs(p, text, size=10.2)


def width_plan(col_count: int) -> list[float]:
    if col_count <= 1:
        return [6.35]
    if col_count == 2:
        return [2.05, 4.30]
    if col_count == 3:
        return [1.45, 2.45, 2.45]
    if col_count == 4:
        return [1.25, 1.55, 1.85, 1.70]
    if col_count == 5:
        return [1.0, 1.15, 1.35, 1.35, 1.50]
    each = 6.35 / col_count
    return [each] * col_count


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows: list[list[str]] = []
    for line in lines:
        cells = [clean_inline(cell.strip()) for cell in line.strip().strip("|").split("|")]
        if all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    for row in rows:
        while len(row) < col_count:
            row.append("")

    widths = width_plan(col_count)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                shade_cell(cell, HEADER_FILL)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 18 or c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=0, after=0, line=1.05)
            run = p.add_run(text)
            set_run_font(run, size=8.2 if col_count >= 5 else 8.8, bold=(r_idx == 0))
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.8)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def set_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "CS-202614 复杂航空载体电磁辐射空域特性测量技术"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(header, before=0, after=0, line=1.0)
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("正式提交稿候选 | Level 1/Level 2 半球面路线 | 第 ")
    set_run_font(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)
    run = footer.add_run(" 页")
    set_run_font(run, size=8.5, color=MUTED)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=80, after=10, line=1.1)
    r = p.add_run("复杂航空载体电磁辐射空域特性测量技术方案报告")
    set_run_font(r, size=24, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=24, line=1.1)
    r = p.add_run("CS-202614 | 电磁空间安全全国重点实验室")
    set_run_font(r, size=13.5, color=MUTED)

    rows = [
        ("测量面选择", "2π 上半球面，13 m 半径，162 个测点"),
        ("核心链路", "CST solver-safe 标准源 / Level 2 full48 / 简化结构遮挡对照"),
        ("关键指标", "Level 1 角域 max NMSE=8.41e-5；Level 2 accuracy=1.000；结构 cross-domain accuracy=1.000"),
        ("交付状态", "正式报告 DOCX/PDF 生成；PPTX/视频仍需按 G5 审计继续导出"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [1.55, 4.80])
    table.style = "Table Grid"
    for row, (label, value) in zip(table.rows, rows):
        shade_cell(row.cells[0], HEADER_FILL)
        for idx, text in enumerate((label, value)):
            p = row.cells[idx].paragraphs[0]
            set_paragraph_spacing(p, before=2, after=2, line=1.15)
            run = p.add_run(text)
            set_run_font(run, size=10.2, bold=(idx == 0))

    add_callout(
        doc,
        "边界声明：Level 2 结构证据为 simplified aircraft occlusion transfer on CST-derived element-library fields，"
        "用于约束安装/遮挡敏感性；不是 full-wave CST airframe scattering。",
    )
    doc.add_page_break()


def heading_style(level: int) -> str:
    if level <= 2:
        return "Heading 1"
    if level == 3:
        return "Heading 2"
    return "Heading 3"


def build_toc(doc: Document, headings: list[tuple[int, str]]) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("目录").bold = True
    for level, title in headings:
        if level > 3:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.24 * max(level - 2, 0))
        set_paragraph_spacing(p, before=0, after=2, line=1.05)
        run = p.add_run(title)
        set_run_font(run, size=9.5 if level == 3 else 10.2, color=BLUE if level == 2 else MUTED, bold=(level == 2))
    doc.add_page_break()


def parse_headings(lines: Iterable[str]) -> list[tuple[int, str]]:
    headings = []
    for line in lines:
        if line.startswith("## "):
            headings.append((2, clean_inline(line[3:].strip())))
        elif line.startswith("### "):
            headings.append((3, clean_inline(line[4:].strip())))
    return headings


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    if not code_lines:
        return
    text = "\n".join(code_lines)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.35])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=2, after=2, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=8.0, font=CODE_FONT, color=RGBColor(40, 40, 40))


def build_docx() -> dict[str, object]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    SUMMARY_JSON.parent.mkdir(parents=True, exist_ok=True)

    lines = REPORT_MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    set_header_footer(doc)
    add_cover(doc)
    build_toc(doc, parse_headings(lines))

    table_lines: list[str] = []
    code_lines: list[str] = []
    in_code = False
    inserted_figures = 0
    missing_figures = 0

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            add_markdown_table(doc, table_lines)
            table_lines = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            add_code_block(doc, code_lines)
            code_lines = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("# "):
            continue
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
                code_lines = []
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        if not line.strip():
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            title = clean_inline(heading_match.group(2).strip())
            p = doc.add_paragraph(title, style=heading_style(level))
            for run in p.runs:
                set_run_font(run, size=16 if level == 2 else 13 if level == 3 else 12, bold=True)
            for caption, path in FIGURES_BY_HEADING.get(title, []):
                if add_figure(doc, caption, path):
                    inserted_figures += 1
                else:
                    missing_figures += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, before=0, after=4, line=1.12)
            add_inline_runs(p, line[2:].strip(), size=10.2)
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            set_paragraph_spacing(p, before=0, after=4, line=1.12)
            add_inline_runs(p, numbered.group(1).strip(), size=10.2)
            continue

        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=6, line=1.18)
        add_inline_runs(p, line, size=10.6)

    flush_table()
    flush_code()

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True
    p = doc.add_paragraph("附录：文件与复现入口", style="Heading 1")
    for run in p.runs:
        set_run_font(run, size=16, bold=True, color=ACCENT)
    appendices = [
        ("评分证据板", "outputs/scorecard/scorecard.md"),
        ("完成度审计", "outputs/completion_audit/completion_audit.md"),
        ("总控 dashboard", "outputs/master_dashboard/master_status_dashboard.md"),
        ("结构遮挡对照", "outputs/cst_structure_comparison/README_cst_structure_comparison.md"),
        ("复现命令", "docs/reproduce_commands.md"),
    ]
    table = doc.add_table(rows=1 + len(appendices), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1.8, 4.55])
    for idx, header in enumerate(["材料", "路径"]):
        cell = table.cell(0, idx)
        shade_cell(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=9.2, bold=True)
    for r_idx, (label, path) in enumerate(appendices, start=1):
        for c_idx, text in enumerate([label, path]):
            p = table.cell(r_idx, c_idx).paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0, line=1.05)
            run = p.add_run(text)
            set_run_font(run, size=8.8, font=CODE_FONT if c_idx == 1 else BODY_FONT)

    doc.save(OUT_DOCX)
    summary = {
        "docx": str(OUT_DOCX.relative_to(ROOT)),
        "source_markdown": str(REPORT_MD.relative_to(ROOT)),
        "style_preset": "narrative_proposal",
        "header_template": "proposal_centerpiece with technical-report metadata table",
        "inserted_figures": inserted_figures,
        "missing_figures": missing_figures,
        "is_final_submission_report": True,
        "requires_visual_render_qa": True,
    }
    SUMMARY_JSON.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    return summary


def main() -> None:
    summary = build_docx()
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
